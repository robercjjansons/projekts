import re #<--data matching bibliotēka
from datetime import datetime
import requests #<-- bibliotēka priekš html
from bs4 import BeautifulSoup #<-- datascraping bibliotēka 
from docx import Document



lapas_url = "https://edition.cnn.com/"

try:
    atbilde = requests.get(lapas_url, timeout= 100) #<-- mēģinam iegūt mājaslapas URL
    atbilde.raise_for_status() #<-- ja nav iespējams iegūt URL, tad error
except Exception as e:
    print(f"kļūda iegūstot mājaslapas informāciju!") 
    exit(1)
lapas_info = BeautifulSoup(atbilde.text, 'html.parser') #<-- šeit pārleik visu uz BeautifulSoup lai varētu darboties ar info
visas_saites = lapas_info.find_all('a', href= True) #<-- savāc visus HTML linkus
rakstu_saites = [] #<-- saglabā visus ar 'href' piederību HTML tagus (Hypertext reference)
redzetas_saites = set() #<-- šeit saglabā visus 'href' tagus, kurus jau esam apskatījušies

for saite in visas_saites: 
    href = saite['href'] #<-- šeit izvelk URL no linka
    teksts = saite.get_text(strip= True) #<-- šeit izvelk nevajadzīgo(atstarpes) no URL
    
    if '/2025' in href and teksts and href not in redzetas_saites: #<-- šeit pārliecinamies, ka raksts ir par 2025.gadu, tajā nav tukšs teksts un šis ir neredzēts links
        redzetas_saites.add(href)
        if not href.startswith("http"): #<-- ja linkam nav http sākums, tad mēs to pāŗveidojam
           href = "https://edition.cnn.com" + href
        rakstu_saites.append((teksts, href)) #<-- pievieno redzēto sarakstam
    if len(rakstu_saites) == 5:
        break

  doc = Document()                                   #<-- izveidojam dokumentu
    sodien = datetime.now().strftime("%Y-%m-%d")     #<-- nosakām šodienas datumu
    docx_filename = f"{sidien}_cnn.docx"             #<-- dokumenta nosaukums
       
    for idx, (virsraksts, saite) in enumerate(rakstu_saites, start=1):    #<--lejupielādējam rakstus
        try:
            raksts = requuests.get(saite, timeout=10)
            raksts.raise_for_status()
        except Exception as e:
            print(f"klūda, iegūstot ziņu: {virsraksts} - {e}")
            continue
        raksts_soup = BeautifulSoup(raksts.text, 'html.parser')
        try virsraksts_teksts = raksts_soup.find('h1').get_text(strip=True)      #<--h1=galvenais virsraksts
        except Exception:
        virsrakstrs_teksts = virsraksts


   datums_teksts = "(nav atrasts)"   #<--tiek atrasts datums
try:
    meta_tag = raksts_soup.find("meta",{"itemprop":"datepulished"}) or \
      raksts_soup.find("meta",{"name":"pubdate"}
    if meta_tag and meta_tag.get("content")::
        datums_teksts = meta_tag["content"].split("T")[0]
except Exception:
    pass
ievads_teksts= ""   #<--tiek iegūts ievads no raksta un saīsināts līdz 150 rakstazīmēm
if len(ievads_teksts)>150:
    saisinajums = ievads_teksts[:ievads_teksts.rfind('',0,150)]
    ievads_teksts = saisinajums + "..."
doc.add_paragraph(f"{idx}. Virsraksts{virsakts_taksts}")    #<--katrs raksts tiek saglabākts dokumentā
doc.add_paragraph(f"Datums:{datums_teksts}")
doc.add_paragraph(f"Ievads:{ievads_teksts or 'Nav apraksta'}")
doc.add_paragraph(f" Saite:{saite}")
doc.add_paragraph("")
try::
    doc.sava(docx_filename)             #<--fails tiek saglabāts
    print(f"Rezultāti saglabāti:{docx_filename}")
except Exception as e:
    print(f"Kūda, saglabājot dokumentu: {e}")

           



















